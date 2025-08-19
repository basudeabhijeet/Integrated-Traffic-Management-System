# Try AI directly in your favorite apps … Use Gemini to generate drafts and refine content, plus get Gemini Advanced with access to Google’s next-gen AI for ₹1,950.00 ₹0 for 1 month
from docx import Document
from datetime import datetime
import pandas as pd 
import numpy as np
from difflib import get_close_matches
import re

ROOT_DIR = "./"


def closeMatches(patterns, word):
    """Find the closest match for a given word from a list of patterns."""
    matc = get_close_matches(word, patterns, n=1)
    return matc[0] if matc else None

def extract_last_four_digits(text):
    """Extract the last four digits from a given text."""
    match = re.search(r'\d{4}$', text)
    return match.group(0).strip() if match else None

def make_doc(result_set):
    """Generate a document and send email based on the matched vehicle registration."""
    df = pd.read_excel('database.xlsx', engine='openpyxl')
    df['Rno'] = df['Rno'].astype(str)  # Ensure all values in 'Rno' column are strings

    # Create a mapping of last 4 digits to full registration numbers
    last_four_map = {r[-4:].strip(): r for r in df['Rno'] if len(r) >= 4}

    for item in result_set:
        fname, detected_number = item[0], item[1]
        last_four_digits = extract_last_four_digits(detected_number)

        if not last_four_digits:
            continue  # Skip if no last 4 digits extracted

        matched_full_number = last_four_map.get(last_four_digits)

        if not matched_full_number:
            continue  # Skip if no match found

        match_row = df[df['Rno'] == matched_full_number]

        if match_row.empty:
            continue  # Skip if no match found in database

        email = match_row['Email'].values[0]
        if pd.isna(email) or email == "":
            continue  # Skip if no email found

        # Proceed with document generation and email sending
        im_path = f'./images/{fname}'
        per_path = f'./person/{fname}'
        now = datetime.now()

        doc = Document()
        doc.add_heading('Issued by Traffic Regulations Authority', 1)

        table = doc.add_table(rows=1, cols=2)
        data = [
            ['Registration Number', matched_full_number],
            ['Name of Owner', match_row['Name'].values[0]],
            ['Riding a motor cycle without helmet', '1000'],
            ['Total Fine Amount', '1000 INR']
        ]

        row = table.rows[0].cells
        row[0].text = 'Specification'
        row[1].text = 'Value'

        for ref, val in data:
            row = table.add_row().cells
            row[0].text = ref
            row[1].text = str(val)

        doc.add_picture(im_path)
        doc.add_picture(per_path)
        save_dir = f'./challan/{fname}.docx'
        doc.save(save_dir)

        from utils import send_email  # Import inside function to avoid circular imports
        send_email(save_dir, email)


